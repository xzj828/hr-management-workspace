from pathlib import Path
from tempfile import TemporaryDirectory

from django.test import SimpleTestCase
from docx import Document
from PIL import Image, ImageDraw
from reportlab.pdfgen import canvas

from recruitment.services.file_extraction import ExtractionError, extract_file


class FakeOcr:
    def __init__(self, text="OCR candidate resume"):
        self.text = text
        self.calls = 0

    def read(self, image):
        self.calls += 1
        self.last_size = image.size
        return [{"text": self.text, "bbox": [10, 20, 200, 60], "confidence": 0.98}]


class FakeDocConverter:
    def convert(self, source, output_dir):
        destination = Path(output_dir) / f"{Path(source).stem}.docx"
        document = Document()
        document.add_paragraph("Converted legacy requirement")
        document.save(destination)
        return destination


class MissingDocConverter:
    def convert(self, source, output_dir):
        raise FileNotFoundError("soffice not found")


class FileExtractionTests(SimpleTestCase):
    def test_docx_keeps_paragraph_and_table_order(self):
        with TemporaryDirectory() as directory:
            path = Path(directory) / "requirement.docx"
            document = Document()
            document.add_heading("Product manager", level=1)
            document.add_paragraph("Five years of platform experience")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Required"
            table.cell(0, 1).text = "B2B experience"
            document.save(path)

            result = extract_file(path, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        self.assertEqual(result.method, "docx")
        self.assertEqual([block["text"] for block in result.blocks], [
            "Product manager",
            "Five years of platform experience",
            "Required | B2B experience",
        ])
        self.assertEqual([block["section"] for block in result.blocks], ["heading", "paragraph", "table"])
        self.assertEqual(len({block["id"] for block in result.blocks}), 3)

    def test_legacy_doc_uses_converter_inside_a_temporary_directory(self):
        with TemporaryDirectory() as directory:
            path = Path(directory) / "legacy.doc"
            path.write_bytes(b"legacy-word-placeholder")

            result = extract_file(path, content_type="application/msword", converter=FakeDocConverter())

        self.assertEqual(result.method, "doc_convert")
        self.assertIn("Converted legacy requirement", result.plain_text)

    def test_legacy_doc_reports_missing_libreoffice(self):
        with TemporaryDirectory() as directory:
            path = Path(directory) / "legacy.doc"
            path.write_bytes(b"legacy-word-placeholder")

            with self.assertRaises(ExtractionError) as caught:
                extract_file(path, content_type="application/msword", converter=MissingDocConverter())

        self.assertEqual(caught.exception.code, "libreoffice_unavailable")

    def test_text_pdf_returns_page_numbered_blocks_without_ocr(self):
        with TemporaryDirectory() as directory:
            path = Path(directory) / "resume.pdf"
            pdf = canvas.Canvas(str(path))
            pdf.drawString(72, 760, "Candidate has five years of product management experience and delivered enterprise platforms.")
            pdf.drawString(72, 730, "Led discovery, roadmap planning, analytics and cross-functional delivery for multiple releases.")
            pdf.save()
            ocr = FakeOcr()

            result = extract_file(path, content_type="application/pdf", ocr=ocr)

        self.assertEqual(result.method, "pdf_text")
        self.assertEqual(ocr.calls, 0)
        self.assertEqual(result.blocks[0]["page"], 1)
        self.assertIn("five years", result.plain_text)

    def test_scanned_pdf_renders_pages_and_falls_back_to_ocr(self):
        with TemporaryDirectory() as directory:
            image_path = Path(directory) / "scan.png"
            image = Image.new("RGB", (500, 300), "white")
            ImageDraw.Draw(image).text((20, 20), "scanned resume", fill="black")
            image.save(image_path)
            path = Path(directory) / "scanned.pdf"
            pdf = canvas.Canvas(str(path), pagesize=(500, 300))
            pdf.drawImage(str(image_path), 0, 0, width=500, height=300)
            pdf.save()
            ocr = FakeOcr("Five years product management")

            result = extract_file(path, content_type="application/pdf", ocr=ocr)

        self.assertEqual(result.method, "pdf_ocr")
        self.assertEqual(ocr.calls, 1)
        self.assertEqual(result.blocks[0]["page"], 1)
        self.assertEqual(result.blocks[0]["bbox"], [10, 20, 200, 60])

    def test_png_uses_ocr_and_keeps_bounding_boxes(self):
        with TemporaryDirectory() as directory:
            path = Path(directory) / "online-resume.png"
            Image.new("RGB", (320, 200), "white").save(path)
            ocr = FakeOcr("Online resume content")

            result = extract_file(path, content_type="image/png", ocr=ocr)

        self.assertEqual(result.method, "image_ocr")
        self.assertEqual(result.blocks[0]["bbox"], [10, 20, 200, 60])
        self.assertEqual(result.plain_text, "Online resume content")
